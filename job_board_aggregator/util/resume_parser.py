"""Resume document parsing utilities."""

import os
import io
import tempfile
from typing import Optional, Tuple, Dict, Any
from pathlib import Path

# Document parsing libraries
try:
    from pdfminer.high_level import extract_text as pdf_extract_text
    from pdfminer.pdfinterp import PDFResourceManager
    from pdfminer.converter import TextConverter
    from pdfminer.layout import LAParams
    from pdfminer.pdfpage import PDFPage
    from pdfminer.pdfinterp import PDFPageInterpreter
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

from job_board_aggregator.config import logger


class ResumeParsingError(Exception):
    """Custom exception for resume parsing errors."""
    pass


class ResumeParser:
    """Document parser for resume files."""
    
    # Supported file extensions
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt'}
    
    # Maximum file size (10MB)
    MAX_FILE_SIZE = 10 * 1024 * 1024
    
    def __init__(self):
        """Initialize the resume parser."""
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check if required parsing libraries are available."""
        if not PDF_AVAILABLE:
            logger.warning("pdfminer.six not available - PDF parsing will fail")
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available - DOCX parsing will fail")
    
    def validate_file(self, file_content: bytes, filename: str) -> Tuple[bool, str]:
        """
        Validate uploaded file.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Tuple of (is_valid, error_message)
        """
        # Check file size
        if len(file_content) > self.MAX_FILE_SIZE:
            return False, f"File size ({len(file_content)} bytes) exceeds maximum allowed size ({self.MAX_FILE_SIZE} bytes)"
        
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in self.SUPPORTED_EXTENSIONS:
            return False, f"Unsupported file type: {file_ext}. Supported types: {', '.join(self.SUPPORTED_EXTENSIONS)}"
        
        # Check if content is not empty
        if len(file_content) == 0:
            return False, "File is empty"
        
        return True, ""
    
    def parse_resume(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse resume from uploaded file.
        
        Args:
            file_content: File content as bytes
            filename: Original filename
            
        Returns:
            Dict containing parsed text and metadata
        """
        # Validate file first
        is_valid, error_msg = self.validate_file(file_content, filename)
        if not is_valid:
            raise ResumeParsingError(error_msg)
        
        file_ext = Path(filename).suffix.lower()
        
        try:
            if file_ext == '.pdf':
                text, method = self._parse_pdf(file_content)
            elif file_ext in ['.docx', '.doc']:
                text, method = self._parse_docx(file_content)
            elif file_ext == '.txt':
                text, method = self._parse_txt(file_content)
            else:
                raise ResumeParsingError(f"Unsupported file type: {file_ext}")
            
            # Clean and process the extracted text
            cleaned_text = self._clean_text(text)
            
            return {
                'text': cleaned_text,
                'original_length': len(text),
                'cleaned_length': len(cleaned_text),
                'parsing_method': method,
                'filename': filename,
                'file_extension': file_ext,
                'success': True
            }
            
        except Exception as e:
            logger.error(f"Error parsing resume {filename}: {e}")
            raise ResumeParsingError(f"Failed to parse document: {str(e)}")
    
    def _parse_pdf(self, file_content: bytes) -> Tuple[str, str]:
        """Parse PDF file content."""
        if not PDF_AVAILABLE:
            raise ResumeParsingError("PDF parsing not available - pdfminer.six not installed")
        
        try:
            # Method 1: Simple extraction
            with io.BytesIO(file_content) as file_obj:
                text = pdf_extract_text(file_obj)
                if text and text.strip():
                    return text, "pdf_simple_extraction"
        except Exception as e:
            logger.warning(f"Simple PDF extraction failed: {e}")
        
        try:
            # Method 2: Advanced extraction with layout analysis
            with io.BytesIO(file_content) as file_obj:
                resource_manager = PDFResourceManager()
                fake_file_handle = io.StringIO()
                converter = TextConverter(resource_manager, fake_file_handle, laparams=LAParams())
                page_interpreter = PDFPageInterpreter(resource_manager, converter)
                
                for page in PDFPage.get_pages(file_obj, caching=True, check_extractable=True):
                    page_interpreter.process_page(page)
                
                text = fake_file_handle.getvalue()
                converter.close()
                fake_file_handle.close()
                
                if text and text.strip():
                    return text, "pdf_advanced_extraction"
                else:
                    raise ResumeParsingError("No text could be extracted from PDF")
        except Exception as e:
            raise ResumeParsingError(f"PDF parsing failed: {str(e)}")
    
    def _parse_docx(self, file_content: bytes) -> Tuple[str, str]:
        """Parse DOCX/DOC file content."""
        if not DOCX_AVAILABLE:
            raise ResumeParsingError("DOCX parsing not available - python-docx not installed")
        
        try:
            with io.BytesIO(file_content) as file_obj:
                document = Document(file_obj)
                
                # Extract text from paragraphs
                paragraphs = []
                for paragraph in document.paragraphs:
                    if paragraph.text.strip():
                        paragraphs.append(paragraph.text.strip())
                
                # Extract text from tables
                for table in document.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                paragraphs.append(cell.text.strip())
                
                text = '\n\n'.join(paragraphs)
                
                if not text.strip():
                    raise ResumeParsingError("No text found in document")
                
                return text, "docx_extraction"
                
        except Exception as e:
            raise ResumeParsingError(f"DOCX parsing failed: {str(e)}")
    
    def _parse_txt(self, file_content: bytes) -> Tuple[str, str]:
        """Parse plain text file content."""
        try:
            # Try UTF-8 first
            text = file_content.decode('utf-8')
            return text, "txt_utf8"
        except UnicodeDecodeError:
            try:
                # Fall back to latin-1
                text = file_content.decode('latin-1')
                return text, "txt_latin1"
            except UnicodeDecodeError:
                try:
                    # Fall back to cp1252 (Windows)
                    text = file_content.decode('cp1252')
                    return text, "txt_cp1252"
                except UnicodeDecodeError as e:
                    raise ResumeParsingError(f"Text file encoding not supported: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text."""
        if not text:
            return ""
        
        # Remove excessive whitespace
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            # Strip whitespace from each line
            cleaned_line = line.strip()
            if cleaned_line:
                cleaned_lines.append(cleaned_line)
        
        # Join lines with single newlines
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Remove multiple consecutive newlines (more than 2)
        import re
        cleaned_text = re.sub(r'\n{3,}', '\n\n', cleaned_text)
        
        # Remove excessive spaces
        cleaned_text = re.sub(r' {2,}', ' ', cleaned_text)
        
        return cleaned_text.strip()


# Create a singleton instance
resume_parser = ResumeParser()


def parse_resume_file(file_content: bytes, filename: str) -> Dict[str, Any]:
    """
    Convenience function to parse a resume file.
    
    Args:
        file_content: File content as bytes
        filename: Original filename
        
    Returns:
        Dict containing parsed text and metadata
    """
    return resume_parser.parse_resume(file_content, filename)
